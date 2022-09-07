import requests
from bs4 import BeautifulSoup
import re
import pdfplumber
import pandas as pd
import docx
from striprtf.striprtf import rtf_to_text
import win32com.client 
import os
import re


HEADERS = {
    "Accept-Language": "en-US,en;q=0.9",
    "User-Agent": "Chrome/87.0.4280.141"
}

def write_to_text(nursery_name,append,text):

    """Helper function to either write to a new file or append to an existing text file"""

    if append:       
        with open(f"./Data/TextFiles/{nursery_name}.txt","a",encoding='utf-8') as f:
            f.write(text)
    else:
        with open(f"./Data/TextFiles/{nursery_name}.txt","w",encoding='utf-8') as f:
            f.write(text)


def get_text_file_html(url,nursery_name,append):

    """Response object -> Soup -> Text.
    """

    #Send an HTTP requests to the url, save content as unicode. This is the full html document. 
    response = requests.get(url = url, headers = HEADERS)
    page = response.text
    page = re.sub("<br>"," <br> ",page) #Add a space between line breaks to avoid cases of names being join together (will not be matched)

    #Store as BS object. Extract all text in the document. 
    #i.e. "<a href="http://example.com/">\nI linked to <i>example.com</i>\n</a" -> '\nI linked to example.com\n' 
    #.text is invokes get_text() with default arguments
    soup = BeautifulSoup(page,features='html.parser')


    text = soup.text
    text = re.sub('\n',' ',text) #For readability in testing
    text = text.lower()


    if nursery_name == "Hungry Hook Bainbridge":
        text = page.lower()


    write_to_text(nursery_name,append,text)
  


def get_text_file_pdf(url,nursery_name,append):
    
    """This function opens the web content as a pdf using the pdfplumber module. Then iterates through the pages, concatenting 
    text from each page. """


    response = requests.get(url = url, headers = HEADERS)

    #Write to a temporary storage, so it can be read in using pdfplumber
    with open('tmp/catalog.pdf', 'wb') as f:
        f.write(response.content)


    with pdfplumber.open('tmp/catalog.pdf') as pdf:
        pages = pdf.pages
        text = ""
        for page in pages:
            text += page.extract_text()

    text = text.lower()

    write_to_text(nursery_name,append,text)


def get_text_file_docx(url,nursery_name,append):

    """Use docx module to read in text as paragraphs. Then, write concatenated string to 
    TextFiles"""

    
    response = requests.get(url = url, headers = HEADERS)

    with open('tmp/catalog.docx', 'wb') as f:
        f.write(response.content)

    doc = docx.Document('tmp/catalog.docx')
    text = ""
    for docpara in doc.paragraphs:
        text += docpara.text
    text = text.lower()

    write_to_text(nursery_name,append,text)




def get_text_file_rtf(url,nursery_name,append):

    """Use the striprtf module to remove rtf embeddings from text"""

    
    response = requests.get(url = url, headers = HEADERS)


    with open('tmp/catalog.rtf', 'wb') as f:
        f.write(response.content)

    with open('tmp/catalog.rtf') as f:
        text = f.read()
        text = rtf_to_text(text,errors="ignore")

    text = text.lower()

    write_to_text(nursery_name,append,text)
 


def get_text_file_doc(url,nursery_name,append):

    "Use windows com client to convert doc to docx. Follow procedure for docx files"

    response = requests.get(url = url, headers = HEADERS)

    
    with open('tmp/catalog.doc', 'wb') as f:
        f.write(response.content)


    #Use COM client to convert to .docx

    word = win32com.client.Dispatch("Word.Application")
    in_file = os.path.abspath('tmp/catalog.doc')
    wb = word.Documents.Open(in_file)
    out_file = os.path.abspath("tmp/catalog.docx")
    wb.SaveAs2(out_file, FileFormat=16) # file format for docx
    wb.Close()


    #Read Docx 
    doc = docx.Document('tmp/catalog.docx')
    text = ""
    for docpara in doc.paragraphs:
        text += docpara.text
    text = text.lower()
        
    write_to_text(nursery_name,append,text)


def get_text_file_excel(url,nursery_name,append):

    """
    This functions grabs and writes the data as a bytes object. The bytes file is 
    opened as an excel file with pandas. Then all the entries are concatenated into 
    a single string. 
    """

    response = requests.get(url,headers=HEADERS)

    with open("tmp/catalog.xlsx","wb") as f:
        f.write(response.content)


    multi_sheet_file = pd.ExcelFile("tmp/catalog.xlsx",engine='openpyxl')
    excel_sheet_names = multi_sheet_file.sheet_names
    
    
    text = ""

    for name in excel_sheet_names:
        df = pd.read_excel(multi_sheet_file,sheet_name=name)

        df.dropna(thresh=3,inplace=True)#Two remove Title and other other non-data rows
        df.columns = ["col" + str(i) for i in range(len(df.columns))] #Unamed columns are creating problems
        df = df.astype(str) #This solves a TypeError

        df_string= ""
        
        #Column -> list -> concatenated string for each column
        for col in df.columns:
            string = ' '.join(df[col].tolist())
            df_string += string
        
        text += df_string
    
    text = text.lower()
    
        

    write_to_text(nursery_name,append,text)








    